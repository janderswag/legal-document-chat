"""Move 2g — the .docx red-flag report. Read the generated document back with
python-docx and prove Sam's rider: the scope caveat paragraph is present, every
row carries its verification status, and red flags (potentially missing / not
confirmed) sort above found rows. No new dependency — same writer as the
deposition digest."""

import io
import sys
import unittest
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

PIPELINE_DIR = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PIPELINE_DIR))
import api  # noqa: E402

client = TestClient(api.app)

RUN = {
    "matter": "demo-matter",
    "reviewed": "2026-07-11T00:00:00+00:00",
    "doc_id": None,
    "doc_types": ["nda"],
    "summary": {"found": 1, "potentially_missing": 1, "not_confirmed": 1, "total": 3},
    "results": [
        {"id": "parties", "name": "Parties", "category": "Formation",
         "status": "found", "value": "Nimbus and Pemberton",
         "citations": [{"filename": "msa.pdf", "page": 2}]},
        {"id": "arbitration", "name": "Arbitration", "category": "Disputes",
         "status": "potentially_missing", "value": "Not located in the documents.",
         "citations": []},
        {"id": "insurance", "name": "Insurance", "category": "Risk",
         "status": "not_confirmed", "value": "Maybe somewhere",
         "citations": []},
    ],
}


class TestReviewDocx(unittest.TestCase):
    def render(self):
        r = client.post("/clauses/review.docx", json=RUN)
        self.assertEqual(r.status_code, 200, r.text)
        self.assertIn("wordprocessingml", r.headers["content-type"])
        self.assertIn("contract-review-demo-matter.docx",
                      r.headers["content-disposition"])
        return Document(io.BytesIO(r.content))

    def test_caveat_and_statuses_present(self):
        doc = self.render()
        text = "\n".join(p.text for p in doc.paragraphs)
        # Sam's non-negotiable scope caveat, verbatim anchor phrases
        self.assertIn("matter's most relevant passages", text)
        self.assertIn("not legal advice", text)
        self.assertIn("not a complete review", text)

        table = doc.tables[0]
        self.assertEqual(table.rows[0].cells[1].text, "Verification status")
        statuses = [row.cells[1].text for row in table.rows[1:]]
        self.assertIn("Potentially missing", statuses)
        self.assertIn("Found (span-verified)", statuses)
        self.assertIn("Not confirmed (spans rejected)", statuses)

    def test_red_flags_sort_first_and_cite_present(self):
        table = self.render().tables[0]
        first = table.rows[1].cells
        self.assertIn("Arbitration", first[0].text)          # missing leads
        found_row = [r for r in table.rows[1:]
                     if "Parties" in r.cells[0].text][0]
        self.assertIn("msa.pdf p.2", found_row.cells[3].text)

    def test_doc_types_scope_line(self):
        doc = self.render()
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("document types: nda", text)


if __name__ == "__main__":
    unittest.main(verbosity=2)
