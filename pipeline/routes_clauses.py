"""Contract Review router — the clause checklist over the matter's KB documents.

GET /clauses/taxonomy serves the curated checklist (so the UI can show the clauses it
will check). POST /clauses/review {matter, doc_id?} runs extract_clauses over the matter
(scoped to the dedicated .lancedb_kb KB store, read-only) and returns the structured
result. Read-only: there are NO action verbs and no document is mutated (D-2). The matter
is validated against the catalog allowlist (D-35) before answering. Each verified,
chunk-derived citation (D-19/D-38) is enriched with its catalog doc_id so the UI can reuse
the existing /kb/highlight page-thumbnail + cited-span surface. We never add a
model-asserted page — displayed citations stay exactly the verifier's output.
"""

import hashlib
import json

from fastapi import APIRouter, HTTPException
from pydantic import BaseModel

import catalog
import jobs
import review_job
import routes_kb  # shared KB store path (monkeypatchable in tests)
from clauses import extract_clauses, load_taxonomy, taxonomy_version

router = APIRouter()


class ReviewRequest(BaseModel):
    matter: str
    doc_id: int | None = None  # optional: narrow the checklist to one document


class ReviewJobRequest(BaseModel):
    matter: str
    doc_id: int | None = None          # single-document review (council Move 2e)
    doc_types: list[str] | None = None  # attorney-designated; skips inapplicable clauses
    questions: list[str] | None = None  # the attorney's own questions (decision #3)


@router.get("/clauses/taxonomy")
def taxonomy():
    """The curated clause checklist (ids/names/categories/questions) — no document data."""
    return {"clauses": load_taxonomy()}


@router.post("/clauses/review")
def review(body: ReviewRequest):
    """Run the clause checklist for ``matter`` over its KB documents (read-only)."""
    if not catalog.get_matter(body.matter):
        raise HTTPException(status_code=400, detail=f"unknown matter: {body.matter!r}")

    try:
        out = extract_clauses(body.matter, doc_id=body.doc_id,
                              db_path=str(routes_kb.KB_DB))
    except ValueError as e:
        # unknown doc_id, or the matter has no indexed chunks -> a clean 400
        raise HTTPException(status_code=400, detail=str(e))

    # Enrich each verified citation with its catalog doc_id (by matter+filename) so the
    # UI can request /kb/highlight/<doc_id>. The displayed page/span stay chunk-derived
    # (D-38); we add no model-asserted data.
    by_name = {d["filename"]: d["id"] for d in catalog.list_documents(body.matter)}
    for row in out["results"]:
        for c in row["citations"]:
            c["doc_id"] = by_name.get(c["filename"])
    return out


@router.post("/clauses/review-jobs")
def submit_review_job(body: ReviewJobRequest):
    """Queue a clause review on the job runner (D-90; the UI's Run review). Returns
    the job row; if a review for the same matter+doc scope is already queued/running,
    returns THAT job with existing=true — a double click never burns the model twice."""
    if not catalog.get_matter(body.matter):
        raise HTTPException(status_code=400, detail=f"unknown matter: {body.matter!r}")
    if body.doc_id is not None:
        doc = catalog.get_document(body.doc_id)
        if not doc or doc.get("matter_slug") != body.matter:
            raise HTTPException(status_code=400,
                                detail=f"document {body.doc_id} is not in matter "
                                       f"{body.matter!r}")
    params = {"matter": body.matter, "doc_id": body.doc_id,
              "doc_types": body.doc_types, "questions": body.questions}
    # The dedupe key covers the FULL scope (doc + type filter + custom questions):
    # a submit with different settings queues its own run (the runner is serial)
    # instead of silently returning an in-flight job with the wrong scope.
    scope = hashlib.sha256(json.dumps(
        [body.doc_id, body.doc_types, body.questions]).encode()).hexdigest()[:12]
    job = jobs.submit(review_job.KIND, params, matter_slug=body.matter,
                      dedupe_key=f"{review_job.KIND}:{body.matter}:{scope}")
    return job


@router.get("/clauses/runs")
def latest_run(matter: str):
    """The matter's most recent FINISHED review (zero-second reopen), with honest
    staleness: stale=true when the matter's documents or the taxonomy changed since
    the run. Never re-runs anything."""
    if not catalog.get_matter(matter):
        raise HTTPException(status_code=400, detail=f"unknown matter: {matter!r}")
    done = catalog.job_list(kind=review_job.KIND, matter_slug=matter, status="done",
                            limit=1)
    active = (catalog.job_list(kind=review_job.KIND, matter_slug=matter,
                               status="running", limit=1) or
              catalog.job_list(kind=review_job.KIND, matter_slug=matter,
                               status="queued", limit=1))
    out = {"run": None, "active_job_id": active[0]["id"] if active else None}
    if done:
        result = done[0].get("result") or {}
        out["run"] = result
        out["run"]["job_id"] = done[0]["id"]
        out["run"]["stale"] = (
            result.get("docs_key") != review_job.docs_key(matter) or
            result.get("taxonomy_version") != taxonomy_version())
    return out


@router.post("/clauses/review.docx")
def review_docx(payload: dict):
    """Render a finished review (the persisted run JSON) to a Word red-flag report.
    Same local python-docx path as the deposition digest — no new dependency, no
    network. Sam's rider (council 2026-07-11, non-negotiable): every export carries
    the per-clause verification status and the retrieval-scope caveat. Nothing is
    added or rewritten — verbatim findings + cites only."""
    import io

    from docx import Document
    from docx.shared import Pt
    from fastapi.responses import Response

    matter = payload.get("matter", "")
    results = payload.get("results") or []
    summary = payload.get("summary") or {}

    status_label = {"found": "Found (span-verified)",
                    "potentially_missing":
                        "Not located (passages checked, not a page-by-page read)",
                    "not_confirmed": "Not confirmed (spans rejected)"}

    def row_status_label(r):
        # D5: a verified absence carries the per-document claim into the report
        if (r.get("status") == "potentially_missing"
                and r.get("verified_scope") == "matter_documents"):
            return "Not located (each document checked individually)"
        return status_label.get(r.get("status"), str(r.get("status", "")))

    doc = Document()
    doc.add_heading(f"Contract Review - {matter}", level=1)
    scope = payload.get("doc_id")
    doc_types = payload.get("doc_types")
    if not isinstance(doc_types, list):
        doc_types = [doc_types] if doc_types else []
    doc.add_paragraph(f"Reviewed {payload.get('reviewed', '')}"
                      + (" - single document" if scope else " - whole matter")
                      + (f" - document types: {', '.join(str(t) for t in doc_types)}"
                         if doc_types else ""))
    # THE CAVEAT (Sam): scope honesty + verification honesty, on every export.
    doc.add_paragraph(
        "Each clause was checked against the matter's most relevant passages - "
        "not a page-by-page read. A 'Not located' row means the clause was not "
        "located in those passages, not that it is absent. Every 'Found' quote was "
        "mechanically "
        "verified against the source text; verify context before use. This is not "
        "legal advice and not a complete review.")
    doc.add_paragraph(
        f"Summary: {summary.get('found', 0)} found - "
        f"{summary.get('potentially_missing', 0)} not located - "
        f"{summary.get('not_confirmed', 0)} not confirmed "
        f"(of {summary.get('total', 0)} checked)")

    # Red flags first: the rows an attorney acts on lead the report.
    order = {"potentially_missing": 0, "not_confirmed": 1, "found": 2}
    rows = sorted(results, key=lambda r: order.get(r.get("status"), 3))

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "Clause", "Verification status"
    hdr[2].text, hdr[3].text = "Finding", "Cite"
    for r in rows:
        if not isinstance(r, dict):
            continue  # defensive: payload is caller-shaped JSON, never trusted
        cells = table.add_row().cells
        cells[0].text = f"{r.get('name', '')} ({r.get('category', '')})"
        cells[1].text = row_status_label(r)
        cells[2].text = str(r.get("value") or "").strip()
        cites = r.get("citations")
        cites = cites if isinstance(cites, list) else []
        cells[3].text = "; ".join(
            f"{c.get('filename', '')} p.{c.get('page', '')}"
            for c in cites if isinstance(c, dict))
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    slug = "".join(ch if ch.isalnum() or ch in "-_" else "-" for ch in matter)
    return Response(
        buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition":
                 f'attachment; filename="contract-review-{slug or "matter"}.docx"'})
